"""
Document text extractor service.
Supports PDF, DOCX, and TXT files.
"""
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_from_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from PDF bytes. Returns (text, page_count)."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts), page_count
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            page_count = len(reader.pages)
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            return "\n\n".join(text_parts), page_count
        except Exception as e2:
            raise ValueError(f"Failed to extract PDF text: {e2}")


def extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to extract DOCX text: {e}")


def extract_from_txt(file_bytes: bytes) -> str:
    """Extract text from TXT bytes, trying common encodings."""
    for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError("Unable to decode text file with known encodings")


def extract_text(file_bytes: bytes, filename: str) -> dict:
    """
    Main extraction entry point.
    Returns dict with keys: text, page_count (optional), word_count
    """
    ext = Path(filename).suffix.lower()
    page_count = None

    if ext == ".pdf":
        text, page_count = extract_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        text = extract_from_docx(file_bytes)
    elif ext == ".txt":
        text = extract_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: PDF, DOCX, TXT")

    word_count = len(text.split())
    logger.info(f"Extracted {word_count} words from {filename}")

    return {
        "text": text,
        "page_count": page_count,
        "word_count": word_count,
    }
